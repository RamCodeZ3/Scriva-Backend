from __future__ import annotations

import asyncio
import re
from io import BytesIO
from urllib.request import urlopen

from application.dtos.export_result import ExportResult
from application.ports.document_exporter_port import DocumentExporterPort
from docx import Document as DocxDocument
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_BREAK,
    WD_COLOR_INDEX,
    WD_TAB_ALIGNMENT,
    WD_TAB_LEADER,
)
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from domain.entities.document import Document
from domain.exceptions import DocumentBuildError
from domain.value_objects.apa_structure import (
    APASectionType,
    normalize_document_styles,
)
from domain.value_objects.document_node import (
    BLOCK_QUOTE,
    BULLETED_LIST,
    HEADING_1,
    HEADING_2,
    HEADING_3,
    HEADING_4,
    HEADING_5,
    IMAGE,
    NUMBERED_LIST,
    PAGE_BREAK,
    PARAGRAPH,
    TABLE,
    TABLE_OF_CONTENTS,
    DocumentNode,
)

from infrastructure.export.pdf_document_exporter_adapter import (
    PdfDocumentExporterAdapter,
)

_PAGE_SIZES_IN = {"letter": (8.5, 11.0), "a4": (8.2677, 11.6929)}
_ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}
_HEADING_STYLE_NAMES = {
    HEADING_1: "Heading1",
    HEADING_2: "Heading2",
    HEADING_3: "Heading3",
    HEADING_4: "Heading4",
    HEADING_5: "Heading5",
}
_HIGHLIGHT_PALETTE = {
    WD_COLOR_INDEX.YELLOW: "FFFF00",
    WD_COLOR_INDEX.BRIGHT_GREEN: "00FF00",
    WD_COLOR_INDEX.TURQUOISE: "00FFFF",
    WD_COLOR_INDEX.PINK: "FF00FF",
    WD_COLOR_INDEX.RED: "FF0000",
    WD_COLOR_INDEX.BLUE: "0000FF",
    WD_COLOR_INDEX.GRAY_25: "BFBFBF",
}
_SECTION_ATTR = "{urn:scriva:document}section-type"


class DocxDocumentExporterAdapter(DocumentExporterPort):
    async def export(self, document: Document) -> ExportResult:
        try:
            docx_bytes = await asyncio.to_thread(self._build_sync, document)
        except Exception as exc:
            raise DocumentBuildError(f"DOCX export failed: {exc}") from exc

        return ExportResult(
            url=None,
            file_bytes=docx_bytes,
            file_name=f"{_safe_filename(document.title)}.docx",
            content_type=(
                "application/vnd.openxmlformats-officedocument"
                ".wordprocessingml.document"
            ),
        )

    def _build_sync(self, document: Document) -> bytes:
        toc_entries = PdfDocumentExporterAdapter().build_toc_entries(document)
        doc_styles = normalize_document_styles(document.document_styles)
        docx = DocxDocument()
        content_width_pt = _setup_page(docx, doc_styles)
        styles = _build_styles(docx, doc_styles)
        _setup_page_numbers(docx, doc_styles)
        _enable_auto_update_fields(docx)

        ctx = {"styles": styles, "content_width_pt": content_width_pt}

        self._build_cover_page(docx, document, ctx)
        # The page transition is owned by the exporter. Older generated
        # documents may also contain one or more page-break nodes in the
        # presentation section; rendering those as well creates a blank page.
        docx.add_page_break()
        self._build_toc_page(docx, document, ctx, toc_entries)

        for section_type in (
            APASectionType.INTRODUCTION,
            APASectionType.BODY,
            APASectionType.CONCLUSION,
        ):
            self._build_section(docx, document, section_type, ctx)

        self._build_references(docx, document, ctx)

        buffer = BytesIO()
        docx.save(buffer)
        return buffer.getvalue()

    def _build_cover_page(self, docx, document: Document, ctx: dict) -> None:
        section = document.get_section(APASectionType.PRESENTATION)
        title_p = docx.add_paragraph(style=ctx["styles"]["TitleCover"])
        _mark_section(title_p._p, APASectionType.PRESENTATION)
        if section is None:
            title_p.add_run(document.title)
            return
        _render_inline(title_p, section.heading.children)
        _apply_block_style(title_p, section.heading.styles)

        for node in section.body_nodes:
            if node.type == PAGE_BREAK:
                # Normalized once by _build_sync after the complete cover.
                continue
            line = docx.add_paragraph(style=ctx["styles"]["CoverLine"])
            _render_inline(line, node.children)
            _apply_block_style(line, node.styles)

    def _build_toc_page(
        self,
        docx,
        document: Document,
        ctx: dict,
        toc_entries: list[tuple[int, str, int]],
    ) -> None:
        index_section = document.get_section(APASectionType.INDEX)
        index_title = index_section.title if index_section else "Índice"

        # "Heading1Plain" carries no outline level, so it can't register
        # itself as a TOC entry — same reasoning as the PDF adapter's
        # "Heading1Plain" vs "Heading1".
        heading = docx.add_paragraph(style=ctx["styles"]["Heading1Plain"])
        _mark_section(heading._p, APASectionType.INDEX)
        if index_section is None:
            heading.add_run(index_title)
        else:
            _render_inline(heading, index_section.heading.children)
            _apply_block_style(heading, index_section.heading.styles)
        # The index is generated structure, not editable body content. Some
        # browser DOCX editors flatten a complex TOC field into plain
        # paragraphs on save. Always rebuilding it prevents those paragraphs
        # from replacing the updateable Word field.
        toc_styles = {}
        if index_section is not None:
            toc_node = next(
                (
                    node
                    for node in index_section.body_nodes
                    if node.type == TABLE_OF_CONTENTS
                ),
                None,
            )
            if toc_node is not None:
                toc_styles = toc_node.styles
        _insert_toc_field(
            docx, toc_entries, ctx["content_width_pt"], toc_styles
        )
        docx.add_page_break()

    def _build_section(
        self,
        docx,
        document: Document,
        section_type: APASectionType,
        ctx: dict,
    ) -> None:
        section = document.get_section(section_type)
        if section is None:
            return

        if section_type is not APASectionType.BODY:
            heading = docx.add_paragraph(style=ctx["styles"]["Heading1"])
            _mark_section(heading._p, section_type)
            _render_inline(heading, section.heading.children)
            _apply_block_style(heading, section.heading.styles)
        for index, node in enumerate(section.body_nodes):
            _render_block(docx, node, ctx)
            if section_type is APASectionType.BODY and index == 0:
                body_elements = list(docx.element.body)
                # python-docx keeps sectPr as the final body child and
                # inserts new paragraphs/tables immediately before it.
                if len(body_elements) >= 2:
                    _mark_section(body_elements[-2], section_type)

    def _build_references(self, docx, document: Document, ctx: dict) -> None:
        conclusion = document.get_section(APASectionType.CONCLUSION)
        if conclusion is None or conclusion.body_nodes[-1].type != PAGE_BREAK:
            docx.add_page_break()
        sources_section = document.get_section(APASectionType.SOURCES)
        title = sources_section.title if sources_section else "References"
        heading = docx.add_paragraph(style=ctx["styles"]["Heading1"])
        _mark_section(heading._p, APASectionType.SOURCES)
        if sources_section is None:
            heading.add_run(title)
        else:
            _render_inline(heading, sources_section.heading.children)
            _apply_block_style(heading, sources_section.heading.styles)

        imported_references = sources_section is not None and any(
            node.metadata.get("docxImported")
            for node in sources_section.body_nodes
        )
        if sources_section is not None and imported_references:
            for node in sources_section.body_nodes:
                if node.type == PAGE_BREAK:
                    continue
                if node.type == PARAGRAPH:
                    paragraph = docx.add_paragraph(
                        style=ctx["styles"]["Reference"]
                    )
                    _render_inline(paragraph, node.children)
                    _apply_block_style(paragraph, node.styles)
                else:
                    _render_block(docx, node, ctx)
            return

        for ref in sorted(
            document.sources, key=lambda r: (r.author or "").lower()
        ):
            p = docx.add_paragraph(style=ctx["styles"]["Reference"])
            p.add_run(ref.to_apa_string())
        if not document.sources:
            p = docx.add_paragraph(style=ctx["styles"]["Body"])
            p.add_run("No sources were provided.")


def _mark_section(element, section_type: APASectionType) -> None:
    """Embed a non-rendered round-trip marker in the OOXML element."""
    element.set(_SECTION_ATTR, section_type.value)


# --- page setup -----------------------------------------------------------


def _setup_page(docx, doc_styles: dict) -> float:
    section = docx.sections[0]
    width_in, height_in = _resolve_page_size(doc_styles)
    landscape = (
        str(doc_styles.get("orientation", "portrait")).lower() == "landscape"
    )
    if landscape:
        width_in, height_in = (
            max(width_in, height_in),
            min(width_in, height_in),
        )
        section.orientation = WD_ORIENT.LANDSCAPE
    else:
        width_in, height_in = (
            min(width_in, height_in),
            max(width_in, height_in),
        )
        section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(width_in)
    section.page_height = Inches(height_in)

    margins = _resolve_margins(doc_styles)
    section.top_margin = Pt(margins["top"])
    section.bottom_margin = Pt(margins["bottom"])
    section.left_margin = Pt(margins["left"])
    section.right_margin = Pt(margins["right"])

    background = doc_styles.get("backgroundColor")
    if background:
        _set_document_background(docx, background)

    return width_in * 72 - margins["left"] - margins["right"]


def _set_document_background(docx, hex_color) -> None:
    hex_clean = str(hex_color).lstrip("#")
    bg = OxmlElement("w:background")
    bg.set(qn("w:color"), hex_clean)
    docx.element.insert(0, bg)


def _setup_page_numbers(docx, doc_styles: dict) -> None:
    if not doc_styles.get("showPageNumbers", True):
        return
    position = doc_styles.get("pageNumberPosition", "top-right")
    section = docx.sections[0]

    if position in ("bottom-center", "bottom-right"):
        section.footer.is_linked_to_previous = False
        container = section.footer
        align = (
            WD_ALIGN_PARAGRAPH.CENTER
            if position == "bottom-center"
            else WD_ALIGN_PARAGRAPH.RIGHT
        )
    else:  # "top-right" (default)
        section.header.is_linked_to_previous = False
        container = section.header
        align = WD_ALIGN_PARAGRAPH.RIGHT

    p = (
        container.paragraphs[0]
        if container.paragraphs
        else container.add_paragraph()
    )
    p.alignment = align
    p.text = ""
    _add_field(p, "PAGE", cached_text="1")


def _enable_auto_update_fields(docx) -> None:
    """Sets settings.xml's <w:updateFields w:val="true"/> so Word recomputes
    the TOC (and any other field) automatically on open, instead of
    requiring the user to right-click > Update Field."""
    settings = docx.settings.element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)


# --- styles -----------------------------------------------------------


def _build_styles(docx, doc_styles: dict) -> dict[str, str]:
    """Creates the custom paragraph styles this adapter renders with, and
    returns a lookup from logical name -> Word style name. Unlike the PDF
    adapter (which builds throwaway ParagraphStyle objects per call), these
    are real named styles registered on the document, since python-docx
    paragraphs are always created *against* a style."""
    font_name = _resolve_font_family(doc_styles.get("fontFamily"))
    base_size = _parse_length_pt(doc_styles.get("fontSize"), default=12) or 12
    line_height = _coerce_float(doc_styles.get("lineHeight"), default=2.0)
    text_color = _parse_color(doc_styles.get("color"), default="000000")

    normal = docx.styles["Normal"]
    normal.font.name = font_name
    normal.font.size = Pt(base_size)
    normal.font.color.rgb = RGBColor.from_string(text_color)
    normal.paragraph_format.line_spacing = line_height

    def make(
        name,
        *,
        base="Normal",
        bold=False,
        italic=False,
        size=None,
        align=None,
        space_before=None,
        space_after=None,
        left_indent=None,
        right_indent=None,
        first_line_indent=None,
        outline_level=None,
    ):
        style = docx.styles.add_style(name, docx.styles["Normal"].type)
        style.base_style = docx.styles[base]
        style.font.bold = bold
        style.font.italic = italic
        if size is not None:
            style.font.size = Pt(size)
        pf = style.paragraph_format
        if align is not None:
            pf.alignment = align
        if space_before is not None:
            pf.space_before = Pt(space_before)
        if space_after is not None:
            pf.space_after = Pt(space_after)
        if left_indent is not None:
            pf.left_indent = Pt(left_indent)
        if right_indent is not None:
            pf.right_indent = Pt(right_indent)
        if first_line_indent is not None:
            pf.first_line_indent = Pt(first_line_indent)
        if outline_level is not None:
            _set_outline_level(style, outline_level)
        return name

    def make_heading(name, **kwargs):
        # Headings should never be the last line on a page with their body
        # text starting on the next one.
        style_name = make(name, **kwargs)
        docx.styles[style_name].paragraph_format.keep_with_next = True
        return style_name

    return {
        "TitleCover": make(
            "APA Title Cover",
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            space_before=180,
            space_after=36,
        ),
        "CoverLine": make(
            "APA Cover Line",
            align=WD_ALIGN_PARAGRAPH.CENTER,
            space_after=6,
        ),
        # Registers as TOC level 1.
        "Heading1": make_heading(
            "APA Heading 1",
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            space_before=12,
            space_after=12,
            outline_level=0,
        ),
        # Same look as Heading1 but deliberately carries NO outline level,
        # so it never shows up in the TOC (used for the cover title / the
        # "Índice" heading itself).
        "Heading1Plain": make_heading(
            "APA Heading 1 Plain",
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            space_before=12,
            space_after=12,
        ),
        # Registers as TOC level 2.
        "Heading2": make_heading(
            "APA Heading 2",
            bold=True,
            align=WD_ALIGN_PARAGRAPH.LEFT,
            space_before=12,
            space_after=6,
            outline_level=1,
        ),
        "Heading3": make_heading(
            "APA Heading 3",
            bold=True,
            italic=True,
            align=WD_ALIGN_PARAGRAPH.LEFT,
            space_before=10,
            space_after=6,
        ),
        "Heading4": make_heading(
            "APA Heading 4",
            bold=True,
            align=WD_ALIGN_PARAGRAPH.LEFT,
            left_indent=36,
            space_before=8,
            space_after=4,
        ),
        "Heading5": make_heading(
            "APA Heading 5",
            bold=True,
            italic=True,
            align=WD_ALIGN_PARAGRAPH.LEFT,
            left_indent=36,
            space_before=8,
            space_after=4,
        ),
        "Body": make(
            "APA Body",
            align=WD_ALIGN_PARAGRAPH.JUSTIFY,
            first_line_indent=36,
        ),
        "BlockQuote": make(
            "APA Block Quote",
            align=WD_ALIGN_PARAGRAPH.JUSTIFY,
            left_indent=36,
            right_indent=36,
            space_before=6,
            space_after=6,
        ),
        "Bullet": "List Bullet",
        "Numbered": "List Number",
        "Reference": make(
            "APA Reference",
            align=WD_ALIGN_PARAGRAPH.LEFT,
            left_indent=36,
            first_line_indent=-36,
            space_after=12,
        ),
        "Caption": make(
            "APA Caption",
            italic=True,
            size=max(base_size - 2, 8),
            align=WD_ALIGN_PARAGRAPH.CENTER,
            space_before=4,
            space_after=12,
        ),
    }


def _set_outline_level(style, level: int) -> None:
    """Word only registers a heading in a `{ TOC \\o "1-2" }` field if the
    paragraph's outline level falls in that range. Built-in Heading1/2
    styles carry this already; custom styles need it set explicitly (see
    module docstring) or the field silently comes back empty."""
    pPr = style.element.get_or_add_pPr()
    outline = OxmlElement("w:outlineLvl")
    outline.set(qn("w:val"), str(level))
    pPr.append(outline)


# --- table of contents / fields -----------------------------------------


def _insert_toc_field(
    docx,
    entries: list[tuple[int, str, int]],
    content_width_pt: float,
    styles: dict | None = None,
) -> None:
    """Insert an updateable TOC field with an already rendered result."""
    if not entries:
        paragraph = docx.add_paragraph()
        _apply_block_style(paragraph, styles or {})
        _add_field(paragraph, 'TOC \\o "1-2" \\h \\z \\u')
        return

    paragraphs = []
    for level, title, page_number in entries:
        paragraph = docx.add_paragraph()
        paragraph.paragraph_format.left_indent = Pt(18 * level)
        paragraph.paragraph_format.tab_stops.add_tab_stop(
            Pt(content_width_pt),
            WD_TAB_ALIGNMENT.RIGHT,
            WD_TAB_LEADER.DOTS,
        )
        _apply_block_style(paragraph, styles or {})
        paragraph.add_run(title)
        paragraph.add_run("\t")
        paragraph.add_run(str(page_number))
        paragraphs.append(paragraph)

    _start_complex_field(paragraphs[0], 'TOC \\o "1-2" \\h \\z \\u')
    _end_complex_field(paragraphs[-1])


def _start_complex_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    paragraph._p.remove(run._r)
    paragraph._p.insert(0, run._r)

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    run._r.append(begin)

    field_instruction = OxmlElement("w:instrText")
    field_instruction.set(qn("xml:space"), "preserve")
    field_instruction.text = f" {instruction} "
    run._r.append(field_instruction)

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run._r.append(separate)


def _end_complex_field(paragraph) -> None:
    run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(end)


def _add_field(paragraph, instruction: str, *, cached_text: str = "") -> None:
    """Inserts a Word field (begin / instrText / separate / cached-result /
    end) into `paragraph`. `cached_text` is only what shows before Word
    recalculates the field — with `updateFields` enabled (see
    `_enable_auto_update_fields`), that happens automatically on open."""
    run = paragraph.add_run()

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instruction} "
    run._r.append(instr)

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run._r.append(separate)

    if cached_text:
        cached = OxmlElement("w:t")
        cached.set(qn("xml:space"), "preserve")
        cached.text = cached_text
        run._r.append(cached)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(end)


# --- inline rendering (marks) ------------------------------------------


def _render_inline(paragraph, nodes: tuple[DocumentNode, ...]) -> None:
    for node in nodes:
        _render_leaf(paragraph, node)


def _render_leaf(paragraph, node: DocumentNode) -> None:
    if node.text is None:
        raise DocumentBuildError(
            f"Expected a leaf text node, got block '{node.type}'."
        )
    by_type = {m.type: m.value for m in node.marks}

    if "link" in by_type:
        url = str(by_type["link"].get("url", ""))
        run = _add_hyperlink(paragraph, node.text, url)
    else:
        run = paragraph.add_run(node.text)

    run.font.bold = "bold" in by_type or run.font.bold
    run.font.italic = "italic" in by_type or run.font.italic
    run.font.underline = "underline" in by_type or run.font.underline
    run.font.strike = "strikethrough" in by_type or run.font.strike

    if "code" in by_type:
        run.font.name = "Courier New"
    if "color" in by_type:
        hex_color = _parse_color(by_type["color"], default=None)
        if hex_color:
            run.font.color.rgb = RGBColor.from_string(hex_color)
    if "fontSize" in by_type:
        size_pt = _parse_length_pt(by_type["fontSize"], default=None)
        if size_pt:
            run.font.size = Pt(size_pt)
    if "fontFamily" in by_type:
        run.font.name = _resolve_font_family(str(by_type["fontFamily"]))
    if "script" in by_type:
        run.font.subscript = by_type["script"] == "subscript"
        run.font.superscript = by_type["script"] == "superscript"
    if "highlight" in by_type:
        run.font.highlight_color = _closest_highlight(by_type["highlight"])


def _add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(color)
    rPr.append(underline)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

    # Wrap it in a lightweight object exposing the same `.font` surface the
    # caller uses, backed by the run we just built.
    from docx.text.run import Run

    return Run(new_run, paragraph)


def _closest_highlight(hex_value) -> WD_COLOR_INDEX:
    target = _parse_color(hex_value, default=None)
    if target is None:
        return WD_COLOR_INDEX.YELLOW
    target_rgb = tuple(int(target[i : i + 2], 16) for i in (0, 2, 4))
    best, best_dist = WD_COLOR_INDEX.YELLOW, float("inf")
    for idx, hex_code in _HIGHLIGHT_PALETTE.items():
        rgb = tuple(int(hex_code[i : i + 2], 16) for i in (0, 2, 4))
        dist = sum((a - b) ** 2 for a, b in zip(target_rgb, rgb))
        if dist < best_dist:
            best, best_dist = idx, dist
    return best


# --- block rendering ------------------------------------------------------


def _render_block(container, node: DocumentNode, ctx: dict) -> None:
    """`container` is anything exposing python-docx's `add_paragraph` /
    `add_table` — a `Document` or a table `_Cell` both qualify, which is
    what lets `_render_table` recurse into cells with this same function
    (mirrors the PDF adapter's `_render_block` reuse for table cells)."""
    styles = ctx["styles"]

    if node.type == PAGE_BREAK:
        p = container.add_paragraph()
        p.add_run().add_break(WD_BREAK.PAGE)
        return

    if node.type in _HEADING_STYLE_NAMES:
        p = container.add_paragraph(
            style=styles[_HEADING_STYLE_NAMES[node.type]]
        )
        _render_inline(p, node.children)
        _apply_block_style(p, node.styles)
        return

    if node.type == PARAGRAPH:
        p = container.add_paragraph(style=styles["Body"])
        _render_inline(p, node.children)
        _apply_block_style(p, node.styles)
        return

    if node.type == BLOCK_QUOTE:
        p = container.add_paragraph(style=styles["BlockQuote"])
        _render_inline(p, node.children)
        _apply_block_style(p, node.styles)
        return

    if node.type == BULLETED_LIST:
        for item in node.children:
            p = container.add_paragraph(style=styles["Bullet"])
            _render_inline(p, item.children)
            _apply_block_style(p, node.styles)
        return

    if node.type == NUMBERED_LIST:
        for item in node.children:
            p = container.add_paragraph(style=styles["Numbered"])
            _render_inline(p, item.children)
            _apply_block_style(p, node.styles)
        return

    if node.type == IMAGE:
        _render_image(container, node, ctx)
        return

    if node.type == TABLE:
        _render_table(container, node, ctx)
        return

    raise DocumentBuildError(
        f"Unsupported block node in section: '{node.type}'"
    )


def _apply_block_style(paragraph, node_styles: dict) -> None:
    if not node_styles:
        return
    pf = paragraph.paragraph_format

    if "textAlign" in node_styles:
        align = _ALIGN_MAP.get(str(node_styles["textAlign"]).lower())
        if align is not None:
            pf.alignment = align
    if "textIndent" in node_styles:
        value = _parse_length_pt(node_styles["textIndent"], default=None)
        if value is not None:
            pf.first_line_indent = Pt(value)
    if "marginTop" in node_styles:
        value = _parse_length_pt(node_styles["marginTop"], default=None)
        if value is not None:
            pf.space_before = Pt(value)
    if "marginBottom" in node_styles:
        value = _parse_length_pt(node_styles["marginBottom"], default=None)
        if value is not None:
            pf.space_after = Pt(value)
    if "marginLeft" in node_styles:
        value = _parse_length_pt(node_styles["marginLeft"], default=None)
        if value is not None:
            pf.left_indent = Pt(value)
    if "marginRight" in node_styles:
        value = _parse_length_pt(node_styles["marginRight"], default=None)
        if value is not None:
            pf.right_indent = Pt(value)
    if "lineHeight" in node_styles:
        factor = _coerce_float(node_styles["lineHeight"], default=None)
        if factor is not None:
            pf.line_spacing = factor

    # Unlike the PDF adapter (which fakes this via a single-cell Table),
    # Word paragraphs support background/border natively.
    bg = node_styles.get("backgroundColor")
    if bg:
        _shade_paragraph(paragraph, bg)
    border_left = node_styles.get("borderLeft")
    if border_left:
        width_pt, color = _parse_border(border_left)
        if color:
            _set_paragraph_left_border(paragraph, width_pt, color)


def _shade_paragraph(paragraph, hex_color) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), str(hex_color).lstrip("#"))
    pPr.append(shd)


def _set_paragraph_left_border(
    paragraph, width_pt: float, hex_color: str
) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(max(int(width_pt * 8), 4)))  # eighths-of-a-point
    left.set(qn("w:space"), "4")
    left.set(qn("w:color"), hex_color)
    borders.append(left)
    pPr.append(borders)


def _render_image(container, node: DocumentNode, ctx: dict) -> None:
    styles = ctx["styles"]
    content_width_pt = ctx["content_width_pt"]

    p = container.add_paragraph()
    align = str(node.styles.get("alignment", "center")).lower()
    p.alignment = _ALIGN_MAP.get(align, WD_ALIGN_PARAGRAPH.CENTER)

    try:
        image_bytes = _fetch_image_bytes(node.src)
    except Exception:
        placeholder = node.alt or node.caption or node.src or "image"
        run = p.add_run(f"[Image unavailable: {placeholder}]")
        run.font.italic = True
        return

    width_pt = _resolve_dimension(
        node.styles.get("width"), content_width_pt, default=content_width_pt
    )
    p.add_run().add_picture(BytesIO(image_bytes), width=Pt(width_pt))

    if node.caption:
        cap = container.add_paragraph(style=styles["Caption"])
        cap.add_run(node.caption)


def _fetch_image_bytes(src: str | None) -> bytes:
    if not src:
        raise DocumentBuildError("Image node has no 'src'.")
    with urlopen(src, timeout=10) as response:  # noqa: S310 - trusted, app-inserted URLs
        return response.read()


def _render_table(container, node: DocumentNode, ctx: dict) -> None:
    rows = node.children  # each is a TABLE_ROW node
    if not rows:
        raise DocumentBuildError("A 'table' node has no rows.")

    n_cols = len(rows[0].children)
    content_width_pt = ctx["content_width_pt"]
    requested_widths = node.styles.get("columnWidths", ())
    if (
        isinstance(requested_widths, (list, tuple))
        and len(requested_widths) == n_cols
    ):
        column_widths = [
            _resolve_dimension(value, content_width_pt, default=None)
            for value in requested_widths
        ]
        if any(width is None for width in column_widths):
            column_widths = []
    else:
        column_widths = []
    if not column_widths:
        table_width = _resolve_dimension(
            node.styles.get("width"),
            content_width_pt,
            default=content_width_pt,
        )
        column_widths = [table_width / n_cols] * n_cols

    table = container.add_table(rows=len(rows), cols=n_cols)
    table_alignment = {
        "left": WD_TABLE_ALIGNMENT.LEFT,
        "center": WD_TABLE_ALIGNMENT.CENTER,
        "right": WD_TABLE_ALIGNMENT.RIGHT,
    }
    table.alignment = table_alignment.get(
        str(node.styles.get("alignment", "center")).lower(),
        WD_TABLE_ALIGNMENT.CENTER,
    )
    table.autofit = False
    _set_table_width(table, sum(column_widths))
    for index, width in enumerate(column_widths):
        table.columns[index].width = Pt(width)
    _set_table_borders(table)
    table_background = node.styles.get("backgroundColor")

    for r, row in enumerate(rows):
        if len(row.children) != n_cols:
            raise DocumentBuildError(
                "Every 'table-row' must have the same number of "
                f"'table-cell' children (expected {n_cols}, got "
                f"{len(row.children)})."
            )
        row_height = _parse_length_pt(row.styles.get("height"), default=None)
        if row_height is not None:
            table.rows[r].height = Pt(row_height)
        for c, cell_node in enumerate(row.children):  # TABLE_CELL nodes
            cell = table.cell(r, c)
            cell_width = _parse_length_pt(
                cell_node.styles.get("width"), default=column_widths[c]
            )
            cell.width = Pt(cell_width or column_widths[c])
            cell_ctx = {
                **ctx,
                "content_width_pt": cell_width or column_widths[c],
            }
            # python-docx always gives a fresh cell one empty paragraph;
            # drop it once we're about to add real content so we don't
            # leave a blank line above every cell's text.
            cell.paragraphs[0].text = ""
            for child in cell_node.children:
                paragraph_count = len(cell.paragraphs)
                _render_block(cell, child, cell_ctx)
                if (
                    child.type == PARAGRAPH
                    and "textIndent" not in child.styles
                ):
                    for paragraph in cell.paragraphs[paragraph_count:]:
                        paragraph.paragraph_format.first_line_indent = Pt(0)
            if len(cell.paragraphs) > 1 and not cell.paragraphs[0].runs:
                cell.paragraphs[0]._p.getparent().remove(cell.paragraphs[0]._p)
            cell_background = cell_node.styles.get(
                "backgroundColor", table_background
            )
            if cell_background:
                _shade_cell(cell, str(cell_background).lstrip("#"))
            elif r == 0 and len(rows) > 1:
                # First row is conventionally the header row.
                _shade_cell(cell, "F5F5F5")


def _set_table_width(table, width_pt: float) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(round(width_pt * 20)))


def _set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    tbl_pr.append(borders)


def _shade_cell(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


# --- small parsing helpers -------------------------------------------------


def _resolve_page_size(doc_styles: dict) -> tuple[float, float]:
    raw = doc_styles.get("pageSize", "letter")
    if isinstance(raw, dict):
        default_w, default_h = _PAGE_SIZES_IN["letter"]
        width = (
            _parse_length_pt(raw.get("width"), default=default_w * 72)
            or default_w * 72
        ) / 72
        height = (
            _parse_length_pt(raw.get("height"), default=default_h * 72)
            or default_h * 72
        ) / 72
        return width, height
    return _PAGE_SIZES_IN.get(str(raw).lower(), _PAGE_SIZES_IN["letter"])


def _resolve_margins(doc_styles: dict) -> dict[str, float]:
    raw = doc_styles.get("pageMargin", "1in")
    default = _parse_length_pt("1in", default=72) or 72
    if isinstance(raw, dict):
        return {
            side: _parse_length_pt(raw.get(side), default=default) or default
            for side in ("top", "bottom", "left", "right")
        }
    value = _parse_length_pt(raw, default=default) or default
    return {"top": value, "bottom": value, "left": value, "right": value}


def _resolve_font_family(name: str | None) -> str:
    if not name:
        return "Times New Roman"
    return name.split(",")[0].strip().strip('"')


def _resolve_dimension(
    value, content_width_pt: float, *, default: float
) -> float:
    if value is None:
        return default
    text = str(value).strip()
    if text.endswith("%"):
        try:
            pct = float(text[:-1]) / 100.0
        except ValueError:
            return default
        return content_width_pt * pct
    return _parse_length_pt(text, default=default) or default


def _parse_length_pt(value, *, default):
    """Parses a CSS-ish length into points (matches the PDF adapter's
    `_parse_length`, but returns points instead of ReportLab's native
    big-points-are-also-points unit, so it's a 1:1 swap)."""
    if value is None:
        return default
    if isinstance(value, (int, float)):
        return float(value)
    text = str(value).strip().lower()
    try:
        if text.endswith("in"):
            return float(text[:-2]) * 72
        if text.endswith("pt"):
            return float(text[:-2])
        if text.endswith("px"):
            return float(text[:-2]) * 0.75  # 96dpi assumption
        if text.endswith("cm"):
            return float(text[:-2]) * 28.3465
        if text.endswith("mm"):
            return float(text[:-2]) * 2.83465
        return float(text)
    except ValueError:
        return default


def _parse_color(value, *, default):
    if not value:
        return default
    text = str(value).strip().lstrip("#")
    if re.fullmatch(r"[0-9a-fA-F]{6}", text):
        return text.upper()
    return default


def _parse_border(value: str) -> tuple[float, str | None]:
    width_pt = 1.0
    color = None
    for part in str(value).split():
        if re.match(r"^[\d.]+(px|pt|in)$", part):
            width_pt = _parse_length_pt(part, default=width_pt) or width_pt
        elif part.startswith("#") or part.isalpha():
            parsed = _parse_color(part, default=None)
            if parsed is not None:
                color = parsed
    return width_pt, color


def _coerce_float(value, *, default):
    try:
        return float(value)
    except (TypeError, ValueError):
        return default


def _safe_filename(title: str) -> str:
    cleaned = re.sub(r"[^\w\-. ]", "", title).strip()
    cleaned = re.sub(r"\s+", "_", cleaned)
    return cleaned[:80] or "document"
