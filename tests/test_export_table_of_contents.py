from __future__ import annotations

import asyncio
import unittest
from dataclasses import replace
from datetime import UTC, datetime
from io import BytesIO
from uuid import uuid4
from zipfile import ZipFile

from docx import Document as ReadDocx
from domain.entities.document import Document, DocumentStatus
from domain.services.table_of_contents_builder import build_index_section
from domain.value_objects.apa_structure import APASection, APASectionType
from domain.value_objects.document_node import (
    HEADING_1,
    HEADING_2,
    MARK_COLOR,
    MARK_HIGHLIGHT,
    MARK_LINK,
    PAGE_BREAK,
    PARAGRAPH,
    TABLE,
    TABLE_CELL,
    TABLE_OF_CONTENTS,
    TABLE_ROW,
    DocumentNode,
    Mark,
    page_break_node,
    text_node,
)
from domain.value_objects.document_type import DocumentType
from domain.value_objects.presentation_info import PresentationInfo
from infrastructure.export.docx_document_exporter_adapter import (
    DocxDocumentExporterAdapter,
)
from infrastructure.export.pdf_document_exporter_adapter import (
    PdfDocumentExporterAdapter,
)
from infrastructure.parsers.docx_document_parser_adapter import (
    DocxDocumentParserAdapter,
)


class ExportTableOfContentsTest(unittest.TestCase):
    def setUp(self) -> None:
        self.document = _document_fixture()

    def test_pdf_layout_exposes_resolved_toc_entries(self) -> None:
        entries = PdfDocumentExporterAdapter().build_toc_entries(self.document)

        self.assertIn("Introducción", _entry_titles(entries))
        self.assertIn("Tema principal", _entry_titles(entries))
        self.assertTrue(all(page_number > 0 for _, _, page_number in entries))
        self.assertIn((1, "Tema principal"), _entry_levels(entries))

    def test_docx_contains_visible_updateable_toc_on_first_render(
        self,
    ) -> None:
        content = DocxDocumentExporterAdapter()._build_sync(self.document)

        with ZipFile(BytesIO(content)) as archive:
            document_xml = archive.read("word/document.xml").decode("utf-8")

        self.assertIn('TOC \\o "1-2"', document_xml)
        self.assertIn('w:dirty="true"', document_xml)
        self.assertNotIn("Actualizar campos", document_xml)
        self.assertEqual(document_xml.count('w:br w:type="page"'), 3)

        parsed = ReadDocx(BytesIO(content))
        visible_text = "\n".join(
            paragraph.text for paragraph in parsed.paragraphs
        )
        self.assertIn("Introducción", visible_text)
        self.assertIn("Tema principal", visible_text)

    def test_rebuilds_toc_when_an_editor_flattens_it(self) -> None:
        self.document.sections = [
            (
                _section(
                    APASectionType.INDEX,
                    "Índice",
                    (_block(PARAGRAPH, "Flattened stale index"),),
                )
                if section.section_type is APASectionType.INDEX
                else section
            )
            for section in self.document.sections
        ]

        content = DocxDocumentExporterAdapter()._build_sync(self.document)

        with ZipFile(BytesIO(content)) as archive:
            document_xml = archive.read("word/document.xml").decode("utf-8")
        self.assertIn('TOC \\o "1-2"', document_xml)
        self.assertNotIn("Flattened stale index", document_xml)

    def test_parser_preserves_structure_and_text_color(self) -> None:
        colored = DocumentNode(
            type=PARAGRAPH,
            children=(
                text_node(
                    "Colored\nlinked text",
                    marks=(
                        Mark(MARK_COLOR, "#FF0000"),
                        Mark(MARK_HIGHLIGHT, "#FFFF00"),
                        Mark(MARK_LINK, {"url": "https://example.com"}),
                    ),
                ),
            ),
            styles={"textAlign": "right", "marginLeft": "18pt"},
        )
        introduction = self.document.get_section(APASectionType.INTRODUCTION)
        assert introduction is not None
        self.document.sections = [
            (
                _section(
                    APASectionType.INTRODUCTION,
                    "Edited introduction title",
                    (colored,),
                )
                if section.section_type is APASectionType.INTRODUCTION
                else section
            )
            for section in self.document.sections
        ]
        content = DocxDocumentExporterAdapter()._build_sync(self.document)

        sections = asyncio.run(DocxDocumentParserAdapter().parse(content))

        presentation = _parsed_section(sections, APASectionType.PRESENTATION)
        index = _parsed_section(sections, APASectionType.INDEX)
        parsed_introduction = _parsed_section(
            sections, APASectionType.INTRODUCTION
        )
        self.assertEqual(
            sum(node.type == PAGE_BREAK for node in presentation.body_nodes),
            1,
        )
        self.assertTrue(
            any(node.type == TABLE_OF_CONTENTS for node in index.body_nodes)
        )
        marks = parsed_introduction.body_nodes[0].children[0].marks
        self.assertIn(Mark(MARK_COLOR, "#FF0000"), marks)
        self.assertIn(Mark(MARK_HIGHLIGHT, "#FFFF00"), marks)
        self.assertIn(Mark(MARK_LINK, {"url": "https://example.com"}), marks)
        self.assertEqual(
            parsed_introduction.body_nodes[0].styles["textAlign"], "right"
        )
        self.assertEqual(
            parsed_introduction.body_nodes[0].styles["marginLeft"], "18pt"
        )
        self.assertEqual(
            parsed_introduction.body_nodes[0].plain_text(),
            "Colored\nlinked text",
        )
        self.assertEqual(
            parsed_introduction.title, "Edited introduction title"
        )

    def test_parser_preserves_heading_one_and_table_layout(self) -> None:
        table = DocumentNode(
            type=TABLE,
            styles={
                "alignment": "left",
                "columnWidths": ["120pt", "180pt"],
            },
            children=(
                DocumentNode(
                    type=TABLE_ROW,
                    styles={"height": "36pt"},
                    children=(
                        DocumentNode(
                            type=TABLE_CELL,
                            styles={
                                "width": "120pt",
                                "backgroundColor": "#00FF00",
                            },
                            children=(_block(PARAGRAPH, "One"),),
                        ),
                        DocumentNode(
                            type=TABLE_CELL,
                            styles={"width": "180pt"},
                            children=(_block(PARAGRAPH, "Two"),),
                        ),
                    ),
                ),
            ),
        )
        self.document.sections = [
            (
                _section(
                    APASectionType.BODY,
                    section.title,
                    (_block(HEADING_1, "Editable main heading"), table),
                )
                if section.section_type is APASectionType.BODY
                else section
            )
            for section in self.document.sections
        ]

        content = DocxDocumentExporterAdapter()._build_sync(self.document)
        sections = asyncio.run(DocxDocumentParserAdapter().parse(content))
        body = _parsed_section(sections, APASectionType.BODY)

        self.assertEqual(body.body_nodes[0].type, HEADING_1)
        self.assertEqual(
            body.body_nodes[0].plain_text(), "Editable main heading"
        )
        parsed_table = body.body_nodes[1]
        self.assertEqual(parsed_table.type, TABLE)
        self.assertEqual(parsed_table.styles["alignment"], "left")
        self.assertEqual(
            parsed_table.styles["columnWidths"], ["120pt", "180pt"]
        )
        parsed_row = parsed_table.children[0]
        self.assertEqual(parsed_row.styles["height"], "36pt")
        self.assertEqual(
            parsed_row.children[0].styles["backgroundColor"], "#00FF00"
        )

    def test_reference_text_can_be_edited_and_reexported(self) -> None:
        initial = DocxDocumentExporterAdapter()._build_sync(self.document)
        sections = asyncio.run(DocxDocumentParserAdapter().parse(initial))
        sources = _parsed_section(sections, APASectionType.SOURCES)
        edited_reference = replace(
            sources.body_nodes[0],
            children=(text_node("Edited reference text"),),
        )
        self.document.sections = [
            (
                APASection(
                    section.section_type,
                    section.heading,
                    (edited_reference, *section.body_nodes[1:]),
                )
                if section.section_type is APASectionType.SOURCES
                else section
            )
            for section in sections
        ]

        updated = DocxDocumentExporterAdapter()._build_sync(self.document)
        visible_text = "\n".join(
            paragraph.text
            for paragraph in ReadDocx(BytesIO(updated)).paragraphs
        )

        self.assertIn("Edited reference text", visible_text)


def _entry_titles(entries: list[tuple[int, str, int]]) -> set[str]:
    return {title for _, title, _ in entries}


def _entry_levels(entries: list[tuple[int, str, int]]) -> set[tuple[int, str]]:
    return {(level, title) for level, title, _ in entries}


def _parsed_section(
    sections: list[APASection], section_type: APASectionType
) -> APASection:
    return next(
        section for section in sections if section.section_type is section_type
    )


def _document_fixture() -> Document:
    paragraph = _block(PARAGRAPH, "Document content " * 80)
    subheading = _block(HEADING_2, "Tema principal")
    sections = [
        _section(
            APASectionType.PRESENTATION,
            "Presentación",
            (_block(PARAGRAPH, "Student"), page_break_node()),
        ),
        build_index_section(),
        _section(
            APASectionType.INTRODUCTION,
            "Introducción",
            (paragraph,),
        ),
        _section(
            APASectionType.BODY,
            "Desarrollo",
            (subheading, paragraph),
        ),
        _section(
            APASectionType.CONCLUSION,
            "Conclusión",
            (paragraph, page_break_node()),
        ),
        _section(
            APASectionType.SOURCES,
            "Referencias",
            (_block(PARAGRAPH, "Source placeholder"),),
        ),
    ]
    now = datetime.now(UTC)
    return Document(
        id=uuid4(),
        user_id=uuid4(),
        title="TOC test",
        document_type=DocumentType.REPORT,
        raw_sources=[],
        presentation=PresentationInfo("Student", "Professor"),
        status=DocumentStatus.DONE,
        sections=sections,
        sources=[],
        created_at=now,
        updated_at=now,
    )


def _section(
    section_type: APASectionType,
    title: str,
    body: tuple[DocumentNode, ...],
) -> APASection:
    return APASection(
        section_type=section_type,
        heading=_block(HEADING_1, title),
        body_nodes=body,
    )


def _block(node_type: str, text: str) -> DocumentNode:
    return DocumentNode(type=node_type, children=(text_node(text),))


if __name__ == "__main__":
    unittest.main()
